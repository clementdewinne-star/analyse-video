from django.shortcuts import render
from django.http import JsonResponse, FileResponse, HttpResponse
from django.views.decorators.csrf import csrf_exempt
from django.conf import settings
from .models import Video, Sequence
from .ai_google import analyse_tactique
import json
import os
import time
import requests # Nécessaire pour télécharger la vidéo avant découpe
from moviepy import VideoFileClip
from docx import Document

def liste_videos(request):
    videos = Video.objects.all()
    return render(request, 'index.html', {'videos': videos})

def generer_word(texte_ia, titre):
    doc = Document()
    doc.add_heading(f'Rapport : {titre}', 0)
    doc.add_paragraph(texte_ia)
    nom = f"Rapport_{int(time.time())}.docx"
    dossier = os.path.join(settings.MEDIA_ROOT, 'rapports')
    os.makedirs(dossier, exist_ok=True)
    doc.save(os.path.join(dossier, nom))
    # Attention : sur Render, ce fichier sera éphémère, mais téléchargeable immédiatement
    return f"/media/rapports/{nom}"

# Fonction utilitaire pour télécharger depuis Cloudinary
def download_temp(url):
    path = f"temp_dl_{int(time.time())}.mp4"
    with requests.get(url, stream=True) as r:
        r.raise_for_status()
        with open(path, 'wb') as f:
            for chunk in r.iter_content(chunk_size=8192): f.write(chunk)
    return path

def analyser_video_entiere(request, video_id):
    try:
        video = Video.objects.get(id=video_id)
        # On passe l'URL Cloudinary
        rapport = analyse_tactique(video.fichier_video.url)
        url_word = generer_word(rapport, video.titre)
        return JsonResponse({'status': 'ok', 'rapport': rapport, 'url_word': url_word})
    except Exception as e: return JsonResponse({'status': 'error', 'message': str(e)})

def analyser_sequence_ia(request, seq_id):
    try:
        seq = Sequence.objects.get(id=seq_id)
        # On passe l'URL Cloudinary
        rapport = analyse_tactique(seq.video.fichier_video.url, seq.temps_debut, seq.temps_fin)
        url_word = generer_word(rapport, seq.label)
        return JsonResponse({'status': 'ok', 'rapport': rapport, 'url_word': url_word})
    except Exception as e: return JsonResponse({'status': 'error', 'message': str(e)})

def telecharger_sequence(request, seq_id):
    path_temp = None
    try:
        seq = Sequence.objects.get(id=seq_id)
        # 1. Télécharger la source depuis Cloudinary
        path_temp = download_temp(seq.video.fichier_video.url)
        
        # 2. Découper
        nom_out = f"{seq.label}.mp4"
        # On écrit direct dans la réponse pour éviter le stockage disque
        # (Code simplifié pour éviter les erreurs de chemin sur Render)
        # Pour l'instant on garde le stockage temporaire
        path_out = "temp_out_" + nom_out
        
        with VideoFileClip(path_temp) as v:
            v.subclipped(seq.temps_debut, seq.temps_fin).write_videofile(path_out, codec="libx264", audio_codec="aac", preset='ultrafast', logger=None)
        
        # 3. Envoyer
        f = open(path_out, 'rb')
        response = FileResponse(f, as_attachment=True, filename=nom_out)
        
        # Nettoyage (Note: sur Windows/Mac le fichier peut être verrouillé, sur Linux ça va)
        # os.remove(path_temp) 
        
        return response
    except Exception as e: return HttpResponse(str(e))

@csrf_exempt 
def ajouter_tag(request):
    # (Code inchangé par rapport à avant)
    if request.method == 'POST':
        try:
            d = json.loads(request.body)
            v = Video.objects.get(id=d.get('video_id'))
            mode = d.get('mode')
            if mode == 'manual':
                t1, t2 = float(d.get('start_time') or 0), float(d.get('end_time') or 0)
            else:
                t = float(d.get('temps') or 0)
                t1, t2 = max(0, t - float(d.get('lag', 5))), t + float(d.get('lead', 5))
            Sequence.objects.create(video=v, label=d.get('label'), temps_debut=round(t1,1), temps_fin=round(t2,1))
            return JsonResponse({'status': 'ok'})
        except Exception as e: return JsonResponse({'status': 'error', 'message': str(e)})
    return JsonResponse({'status': 'error'})

